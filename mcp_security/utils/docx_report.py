"""
DOCX Report Generator for MCP Security Scans
"""

import logging
from typing import Dict, Any, List
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. DOCX report generation will not be available.")

logger = logging.getLogger(__name__)


class DocxReportGenerator:
    """Generate professional DOCX reports for security scans"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx package is required for DOCX report generation")
    
    def generate_report(self, results: Dict[str, Any], output_path: str):
        """
        Generate a comprehensive DOCX report from scan results
        
        Args:
            results: Dictionary containing scan results
            output_path: Path where the DOCX report should be saved
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx package is required for DOCX report generation")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('MCP Security Scan Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add scan metadata
        doc.add_heading('Scan Information', level=1)
        self._add_scan_info(doc, results)
        
        # Add executive summary
        doc.add_heading('Executive Summary', level=1)
        self._add_executive_summary(doc, results)
        
        # Add prompt injection results
        if 'prompt_injection' in results:
            doc.add_page_break()
            doc.add_heading('Prompt Injection Test Results', level=1)
            self._add_prompt_injection_results(doc, results['prompt_injection'])
        
        # Add penetration test results
        if 'penetration_testing' in results:
            doc.add_page_break()
            doc.add_heading('Penetration Test Results', level=1)
            self._add_pentest_results(doc, results['penetration_testing'])
        
        # Add detailed findings
        doc.add_page_break()
        doc.add_heading('Detailed Findings', level=1)
        self._add_detailed_findings(doc, results)
        
        # Add recommendations
        doc.add_page_break()
        doc.add_heading('Recommendations', level=1)
        self._add_recommendations(doc, results)
        
        # Save the document
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))
        
        logger.info(f"DOCX report saved to {output_file}")
    
    def _add_scan_info(self, doc: Document, results: Dict[str, Any]):
        """Add scan information section"""
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Scan ID
        table.cell(0, 0).text = 'Scan ID'
        table.cell(0, 1).text = results.get('scan_id', 'N/A')
        
        # Target
        table.cell(1, 0).text = 'Target URL'
        table.cell(1, 1).text = results.get('target', 'N/A')
        
        # Timestamp
        table.cell(2, 0).text = 'Timestamp'
        table.cell(2, 1).text = results.get('timestamp', 'N/A')
        
        # Tools Discovered
        table.cell(3, 0).text = 'Tools Discovered'
        table.cell(3, 1).text = str(results.get('tools_discovered', 0))
        
        doc.add_paragraph()
    
    def _add_executive_summary(self, doc: Document, results: Dict[str, Any]):
        """Add executive summary section"""
        summary = results.get('summary', {})
        
        # Risk level with color coding
        risk_level = summary.get('risk_level', 'UNKNOWN')
        para = doc.add_paragraph()
        para.add_run('Risk Level: ').bold = True
        
        risk_run = para.add_run(risk_level)
        risk_run.bold = True
        if risk_level == 'CRITICAL':
            risk_run.font.color.rgb = RGBColor(220, 20, 60)  # Crimson
        elif risk_level == 'HIGH':
            risk_run.font.color.rgb = RGBColor(255, 69, 0)  # Orange Red
        elif risk_level == 'MEDIUM':
            risk_run.font.color.rgb = RGBColor(255, 140, 0)  # Dark Orange
        elif risk_level == 'LOW':
            risk_run.font.color.rgb = RGBColor(255, 215, 0)  # Gold
        else:
            risk_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        
        # Summary table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = 'Total Vulnerabilities'
        table.cell(0, 1).text = str(summary.get('total_vulnerabilities', 0))
        
        table.cell(1, 0).text = 'Critical Vulnerabilities'
        table.cell(1, 1).text = str(summary.get('critical_vulnerabilities', 0))
        
        table.cell(2, 0).text = 'High Vulnerabilities'
        table.cell(2, 1).text = str(summary.get('high_vulnerabilities', 0))
        
        table.cell(3, 0).text = 'Tests Run'
        table.cell(3, 1).text = ', '.join(summary.get('tests_run', []))
        
        doc.add_paragraph()
    
    def _add_prompt_injection_results(self, doc: Document, pi_results: Dict[str, Any]):
        """Add prompt injection test results"""
        summary = pi_results.get('summary', {})
        
        # Summary paragraph
        para = doc.add_paragraph()
        para.add_run(f"Total Tests: {summary.get('total_tests', 0)} | ")
        para.add_run(f"Blocked: {summary.get('injections_blocked', 0)} | ")
        para.add_run(f"Succeeded: {summary.get('injections_succeeded', 0)} | ")
        para.add_run(f"Block Rate: {summary.get('block_rate', 'N/A')}")
        
        doc.add_paragraph()
        
        # LLM Exploitation Summary
        if 'llm_exploitation' in summary:
            llm_summary = summary['llm_exploitation']
            doc.add_heading('LLM Exploitation Analysis', level=2)
            
            para = doc.add_paragraph()
            para.add_run('LLM-Level Exploits: ').bold = True
            exploit_count = llm_summary.get('total_exploited', 0)
            exploit_run = para.add_run(str(exploit_count))
            if exploit_count > 0:
                exploit_run.font.color.rgb = RGBColor(220, 20, 60)  # Red
                exploit_run.bold = True
            
            para.add_run(f" ({llm_summary.get('exploitation_rate', 'N/A')})")
            
            doc.add_paragraph(f"Average Confidence: {llm_summary.get('avg_confidence', 'N/A')}")
            doc.add_paragraph()
        
        # Results by severity
        doc.add_heading('Results by Severity', level=2)
        by_severity = pi_results.get('by_severity', {})
        
        for severity in ['critical', 'high', 'medium', 'low']:
            if severity in by_severity:
                data = by_severity[severity]
                doc.add_heading(severity.upper(), level=3)
                
                para = doc.add_paragraph()
                para.add_run(f"Total: {data.get('total', 0)} | ")
                para.add_run(f"Blocked: {data.get('blocked', 0)} | ")
                
                succeeded_run = para.add_run(f"Succeeded: {data.get('succeeded', 0)}")
                if data.get('succeeded', 0) > 0:
                    succeeded_run.font.color.rgb = RGBColor(220, 20, 60)
                    succeeded_run.bold = True
        
        doc.add_paragraph()
        
        # Detailed test results
        doc.add_heading('Detailed Test Results', level=2)
        
        results_list = pi_results.get('results', [])
        if results_list:
            # Group by tool
            tools_tested = {}
            for result in results_list:
                tool_name = result.get('tool_name', 'Unknown')
                if tool_name not in tools_tested:
                    tools_tested[tool_name] = []
                tools_tested[tool_name].append(result)
            
            for tool_name, tool_results in tools_tested.items():
                doc.add_heading(f'Tool: {tool_name}', level=3)
                
                for result in tool_results[:10]:  # Limit to first 10 per tool
                    para = doc.add_paragraph(style='List Bullet')
                    
                    payload_name = result.get('payload_name', 'Unknown')
                    param = result.get('parameter', 'Unknown')
                    detected = result.get('detected', False)
                    llm_exploited = result.get('llm_exploited', False)
                    
                    para.add_run(f"{payload_name} ").bold = True
                    para.add_run(f"(Parameter: {param}) - ")
                    
                    if llm_exploited:
                        status_run = para.add_run('LLM EXPLOITED')
                        status_run.font.color.rgb = RGBColor(220, 20, 60)
                        status_run.bold = True
                        
                        # Add LLM exploitation details
                        confidence = result.get('confidence', 0)
                        exploit_details = result.get('llm_exploit_details', '')
                        para.add_run(f" (Confidence: {confidence:.0%})")
                        if exploit_details:
                            doc.add_paragraph(f"   └─ {exploit_details}", style='List Bullet 2')
                    elif detected:
                        status_run = para.add_run('BLOCKED')
                        status_run.font.color.rgb = RGBColor(0, 128, 0)
                    else:
                        status_run = para.add_run('BYPASSED')
                        status_run.font.color.rgb = RGBColor(255, 140, 0)
    
    def _add_pentest_results(self, doc: Document, pt_results: Dict[str, Any]):
        """Add penetration test results"""
        summary = pt_results.get('summary', {})
        
        # Summary paragraph
        para = doc.add_paragraph()
        para.add_run(f"Total Tests: {summary.get('total_tests', 0)} | ")
        para.add_run(f"Vulnerabilities Found: {summary.get('vulnerabilities_found', 0)} | ")
        para.add_run(f"Tests Passed: {summary.get('tests_passed', 0)} | ")
        para.add_run(f"Security Score: {summary.get('security_score', 'N/A')}")
        
        doc.add_paragraph()
        
        # Results by severity
        doc.add_heading('Results by Severity', level=2)
        by_severity = pt_results.get('by_severity', {})
        
        for severity in ['critical', 'high', 'medium', 'low']:
            if severity in by_severity:
                data = by_severity[severity]
                doc.add_heading(severity.upper(), level=3)
                
                para = doc.add_paragraph()
                para.add_run(f"Total: {data.get('total', 0)} | ")
                
                vuln_run = para.add_run(f"Vulnerable: {data.get('vulnerable', 0)}")
                if data.get('vulnerable', 0) > 0:
                    vuln_run.font.color.rgb = RGBColor(220, 20, 60)
                    vuln_run.bold = True
                
                para.add_run(f" | Passed: {data.get('passed', 0)}")
        
        doc.add_paragraph()
        
        # Detailed test results
        doc.add_heading('Detailed Test Results', level=2)
        
        results_list = pt_results.get('results', [])
        for result in results_list:
            test_name = result.get('test_name', 'Unknown')
            vulnerable = result.get('vulnerable', False)
            severity = result.get('severity', 'Unknown')
            details = result.get('details', 'No details available')
            
            doc.add_heading(f'{test_name} ({severity})', level=3)
            
            para = doc.add_paragraph()
            para.add_run('Status: ').bold = True
            
            if vulnerable:
                status_run = para.add_run('VULNERABLE')
                status_run.font.color.rgb = RGBColor(220, 20, 60)
                status_run.bold = True
            else:
                status_run = para.add_run('PASSED')
                status_run.font.color.rgb = RGBColor(0, 128, 0)
                status_run.bold = True
            
            doc.add_paragraph(f"Details: {details}")
            
            # Add evidence if available
            evidence = result.get('evidence', {})
            if evidence:
                doc.add_paragraph('Evidence:', style='List Bullet')
                for key, value in evidence.items():
                    if isinstance(value, (str, int, float, bool)):
                        doc.add_paragraph(f"{key}: {value}", style='List Bullet 2')
            
            doc.add_paragraph()
    
    def _add_detailed_findings(self, doc: Document, results: Dict[str, Any]):
        """Add detailed findings section"""
        vulnerabilities = []
        
        # Collect prompt injection vulnerabilities
        if 'prompt_injection' in results:
            pi_results = results['prompt_injection'].get('results', [])
            for result in pi_results:
                if result.get('llm_exploited', False) or not result.get('detected', True):
                    vulnerabilities.append({
                        'type': 'Prompt Injection',
                        'severity': result.get('severity', 'Unknown'),
                        'tool': result.get('tool_name', 'Unknown'),
                        'parameter': result.get('parameter', 'Unknown'),
                        'payload': result.get('payload_name', 'Unknown'),
                        'exploited': result.get('llm_exploited', False),
                        'details': result.get('llm_exploit_details', result.get('indicators', 'No details'))
                    })
        
        # Collect penetration test vulnerabilities
        if 'penetration_testing' in results:
            pt_results = results['penetration_testing'].get('results', [])
            for result in pt_results:
                if result.get('vulnerable', False):
                    vulnerabilities.append({
                        'type': 'Penetration Test',
                        'severity': result.get('severity', 'Unknown'),
                        'test': result.get('test_name', 'Unknown'),
                        'details': result.get('details', 'No details'),
                        'impact': result.get('impact', 'Unknown'),
                        'remediation': result.get('remediation', 'No remediation provided')
                    })
        
        if vulnerabilities:
            # Sort by severity
            severity_order = {'CRITICAL': 0, 'HIGH': 1, 'MEDIUM': 2, 'LOW': 3}
            vulnerabilities.sort(key=lambda x: severity_order.get(x['severity'], 4))
            
            for i, vuln in enumerate(vulnerabilities, 1):
                doc.add_heading(f"Finding {i}: {vuln.get('type')} - {vuln.get('severity')}", level=2)
                
                if vuln['type'] == 'Prompt Injection':
                    doc.add_paragraph(f"Tool: {vuln.get('tool')}")
                    doc.add_paragraph(f"Parameter: {vuln.get('parameter')}")
                    doc.add_paragraph(f"Payload: {vuln.get('payload')}")
                    if vuln.get('exploited'):
                        para = doc.add_paragraph()
                        para.add_run('LLM Exploitation: ').bold = True
                        exploit_run = para.add_run('YES')
                        exploit_run.font.color.rgb = RGBColor(220, 20, 60)
                        exploit_run.bold = True
                    doc.add_paragraph(f"Details: {vuln.get('details')}")
                else:
                    doc.add_paragraph(f"Test: {vuln.get('test')}")
                    doc.add_paragraph(f"Details: {vuln.get('details')}")
                    doc.add_paragraph(f"Impact: {vuln.get('impact')}")
                    doc.add_paragraph(f"Remediation: {vuln.get('remediation')}")
                
                doc.add_paragraph()
        else:
            doc.add_paragraph('No vulnerabilities found. The target system passed all security tests.')
    
    def _add_recommendations(self, doc: Document, results: Dict[str, Any]):
        """Add recommendations section"""
        summary = results.get('summary', {})
        risk_level = summary.get('risk_level', 'UNKNOWN')
        
        doc.add_paragraph(
            'Based on the security scan results, the following recommendations are provided:'
        )
        
        if risk_level in ['CRITICAL', 'HIGH']:
            doc.add_heading('Immediate Actions Required', level=2)
            doc.add_paragraph(
                '1. Address all critical and high-severity vulnerabilities immediately',
                style='List Number'
            )
            doc.add_paragraph(
                '2. Implement input validation and sanitization for all user inputs',
                style='List Number'
            )
            doc.add_paragraph(
                '3. Review and strengthen authentication and authorization mechanisms',
                style='List Number'
            )
            doc.add_paragraph(
                '4. Implement rate limiting and resource quotas',
                style='List Number'
            )
            doc.add_paragraph(
                '5. Conduct a comprehensive security audit',
                style='List Number'
            )
        
        doc.add_heading('General Security Best Practices', level=2)
        doc.add_paragraph(
            '1. Implement comprehensive input validation and output encoding',
            style='List Bullet'
        )
        doc.add_paragraph(
            '2. Use parameterized queries to prevent SQL injection',
            style='List Bullet'
        )
        doc.add_paragraph(
            '3. Implement proper authentication and authorization controls',
            style='List Bullet'
        )
        doc.add_paragraph(
            '4. Use secure defaults and principle of least privilege',
            style='List Bullet'
        )
        doc.add_paragraph(
            '5. Keep all software and dependencies up to date',
            style='List Bullet'
        )
        doc.add_paragraph(
            '6. Implement proper error handling and logging',
            style='List Bullet'
        )
        doc.add_paragraph(
            '7. Use HTTPS for all communications',
            style='List Bullet'
        )
        doc.add_paragraph(
            '8. Implement rate limiting and DDoS protection',
            style='List Bullet'
        )
        doc.add_paragraph(
            '9. Regular security assessments and penetration testing',
            style='List Bullet'
        )
        doc.add_paragraph(
            '10. Security awareness training for development team',
            style='List Bullet'
        )
        
        # LLM-specific recommendations if prompt injection was tested
        if 'prompt_injection' in results:
            doc.add_heading('LLM Security Recommendations', level=2)
            doc.add_paragraph(
                '1. Implement context separation between system prompts and user inputs',
                style='List Bullet'
            )
            doc.add_paragraph(
                '2. Use output sanitization to remove potential prompt injection attempts',
                style='List Bullet'
            )
            doc.add_paragraph(
                '3. Implement semantic validation of LLM outputs',
                style='List Bullet'
            )
            doc.add_paragraph(
                '4. Use LLM-level input filters and guardrails',
                style='List Bullet'
            )
            doc.add_paragraph(
                '5. Monitor LLM behavior for anomalies and unexpected outputs',
                style='List Bullet'
            )
